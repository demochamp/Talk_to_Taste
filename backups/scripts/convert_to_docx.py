import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_content = []
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                continue
            else:
                p = doc.add_paragraph()
                p.style = doc.styles['No Spacing']
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
                continue
        
        if in_code_block:
            code_content.append(line.rstrip())
            continue

        # Handle Tables
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            if '---' in stripped:
                continue
                
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
            continue
        elif in_table:
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_cells in enumerate(table_data):
                    for j, cell_text in enumerate(row_cells):
                        if j < len(table.columns):
                            table.cell(i, j).text = cell_text
            in_table = False
            table_data = []

        # Handle Headings
        if stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        elif not stripped:
            continue
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    
    doc.save(docx_path)

import sys

if __name__ == "__main__":
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace('.md', '.docx')
    else:
        # Default fallback
        md_file = r'C:\Users\masti\.gemini\antigravity\brain\5f059d22-c35b-44dd-8b0f-e9ac42862c70\chapter_6_results_and_discussion.md'
        docx_file = r'c:\Users\masti\Desktop\talkto-taste-voice-assistant\documents\Chapter_6_Results_and_Discussion.docx'
    
    markdown_to_docx(md_file, docx_file)
    print(f"Successfully created {os.path.abspath(docx_file)}")
